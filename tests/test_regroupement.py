import pytest
import importlib.util
from pathlib import Path
from docx import Document

spec = importlib.util.spec_from_file_location(
    "importer", Path(__file__).resolve().parents[1] / "ia_provider" / "importer.py"
)
importer = importlib.util.module_from_spec(spec)
spec.loader.exec_module(importer)
_regrouper_blocs_en_chunks = importer._regrouper_blocs_en_chunks
_est_un_titre_probable = importer._est_un_titre_probable


def _bloc(n):
    return {"niveau_titre": 1, "contenu_markdown": "x" * n, "nombre_mots": n}


def test_un_seul_chunk_si_total_inferieur():
    blocs = [_bloc(1000), _bloc(2000)]
    chunks = _regrouper_blocs_en_chunks(blocs, 7500)
    assert len(chunks) == 1


def test_decoupage_simple():
    blocs = [_bloc(4000), _bloc(3000), _bloc(4000)]
    chunks = _regrouper_blocs_en_chunks(blocs, 7500)
    assert len(chunks) == 2
    assert chunks[0].count("x") == 7000
    assert chunks[1].count("x") == 4000


def test_bloc_geant_isole():
    blocs = [_bloc(1000), _bloc(10000), _bloc(1000)]
    chunks = _regrouper_blocs_en_chunks(blocs, 7500)
    assert len(chunks) == 3
    assert chunks[1].count("x") == 10000


def test_paragraphe_court_gras_est_titre_niveau_2():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Titre bref")
    run.bold = True
    assert _est_un_titre_probable(p) == 2


def test_paragraphe_numerote_gras_est_titre_niveau_1():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("2. Introduction")
    run.bold = True
    assert _est_un_titre_probable(p) == 1


def test_paragraphe_gras_termine_point_pas_titre():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Titre.")
    run.bold = True
    assert _est_un_titre_probable(p) is None


def test_paragraphe_long_gras_pas_titre():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("mots " * 20)
    run.bold = True
    assert _est_un_titre_probable(p) is None


def test_paragraphe_partiel_gras_pas_titre():
    doc = Document()
    p = doc.add_paragraph()
    run1 = p.add_run("Titre")
    run1.bold = True
    p.add_run(" normal")
    assert _est_un_titre_probable(p) is None
