from __future__ import annotations
import io
import logging
from typing import Any, Dict, List

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.text.run import Run
import markdown as md
from bs4 import BeautifulSoup
from bs4.element import NavigableString

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")


class MarkdownToDocxConverter:
    """Convertit du texte Markdown en éléments DOCX."""

    def __init__(self, document: Document, styles: Dict[str, Dict[str, Any]]):
        """Initialise le convertisseur avec un document et un dictionnaire de styles."""

        self.doc = document
        self.styles = styles or {}

    def _apply_style(
        self,
        run,
        style_overrides: Dict[str, Any] | None = None,
        *,
        style_name: str = "response",
        is_heading: bool = False,
    ) -> None:
        """Applique un style au ``run`` donné, en tenant compte du contexte (titre ou non)."""

        style = {**self.styles.get(style_name, {}), **(style_overrides or {})}

        # Si le run appartient à un titre, ne modifier que le gras/italique afin de
        # préserver les attributs du style de titre natif de Word.
        if is_heading:
            run.bold = style.get("is_bold", False)
            run.italic = style.get("is_italic", False)
            return

        # Comportement normal pour le texte hors titres : appliquer l'ensemble
        # des attributs définis.
        if font_name := style.get("font_name"):
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if size := style.get("font_size"):
            run.font.size = Pt(size)
        if color := style.get("font_color_rgb"):
            try:
                if isinstance(color, str):
                    run.font.color.rgb = RGBColor.from_string(color)
                else:
                    run.font.color.rgb = RGBColor(*color)
            except Exception:
                pass
        run.bold = style.get("is_bold", False)
        run.italic = style.get("is_italic", False)

    def _add_inline(self, paragraph, node) -> None:
        """Ajoute récursivement les noeuds inline à un paragraphe."""

        # Détection du contexte : déterminer si le paragraphe utilise un style de
        # titre (nom commençant par "heading" ou "titre").
        is_heading = paragraph.style.name.lower().startswith(("heading", "titre"))

        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(text)
                self._apply_style(run, is_heading=is_heading)
            return

        for child in node.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = paragraph.add_run(text)
                    self._apply_style(run, is_heading=is_heading)
            elif child.name in {"strong", "b"}:
                run = paragraph.add_run(child.get_text())
                self._apply_style(run, {"is_bold": True}, is_heading=is_heading)
            elif child.name in {"em", "i"}:
                run = paragraph.add_run(child.get_text())
                self._apply_style(run, {"is_italic": True}, is_heading=is_heading)
            elif child.name == "code":
                run = paragraph.add_run(child.get_text())
                # Le style par défaut ne doit pas écraser celui d'un titre.
                if not is_heading:
                    self._apply_style(run, is_heading=is_heading)
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            elif child.name == "a":
                text = child.get_text()
                href = child.get("href")
                if href:
                    run = self._add_hyperlink(paragraph, href, text)
                    # Les hyperliens ont leur propre style; éviter de modifier
                    # celui d'un titre.
                    if not is_heading:
                        self._apply_style(run, is_heading=is_heading)
                else:
                    run = paragraph.add_run(text)
                    self._apply_style(run, is_heading=is_heading)
            else:
                self._add_inline(paragraph, child)

    def _add_hyperlink(self, paragraph, url: str, text: str) -> Run:
        """Ajoute un hyperlien cliquable au paragraphe."""

        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        r_pr = OxmlElement('w:rPr')
        r_style = OxmlElement('w:rStyle')
        r_style.set(qn('w:val'), 'Hyperlink')
        r_pr.append(r_style)
        new_run.append(r_pr)

        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

        return Run(new_run, paragraph)

    def _process_element(self, elem, list_style: str | None = None) -> None:
        """Traite les éléments HTML convertis depuis le Markdown."""

        if isinstance(elem, NavigableString):
            text = str(elem).strip()
            if text:
                paragraph = (
                    self.doc.add_paragraph(style=list_style)
                    if list_style
                    else self.doc.add_paragraph()
                )
                run = paragraph.add_run(text)
                self._apply_style(run)
            return

        tag = elem.name
        if tag in {"p", "li"}:
            paragraph = (
                self.doc.add_paragraph(style=list_style)
                if list_style
                else self.doc.add_paragraph()
            )
            self._add_inline(paragraph, elem)
            for child in elem.find_all(["ul", "ol"], recursive=False):
                self._process_element(
                    child,
                    "List Bullet" if child.name == "ul" else "List Number",
                )
        elif tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(tag[1])

            # Utiliser directement la méthode ``add_heading`` afin d'appliquer
            # les styles natifs de Word ("Heading 1", "Heading 2", ...). Cela
            # garantit que la structure du document est correctement
            # interprétée par les lecteurs DOCX et permet d'utiliser les
            # fonctionnalités de navigation de Word.
            paragraph = self.doc.add_heading(level=level)

            # ``add_heading`` insère par défaut un run contenant le texte du
            # titre passé en argument. Comme nous souhaitons gérer nous-même
            # le contenu et les styles inline (gras, italique, liens, ...),
            # nous nettoyons le paragraphe avant de le remplir.
            paragraph.clear()

            # Peupler le paragraphe avec le contenu de ``elem`` en conservant
            # les styles inline éventuels.
            self._add_inline(paragraph, elem)
        elif tag == "ul":
            for li in elem.find_all("li", recursive=False):
                self._process_element(li, "List Bullet")
        elif tag == "ol":
            for li in elem.find_all("li", recursive=False):
                self._process_element(li, "List Number")
        elif tag == "pre":
            code_text = "".join(elem.strings).strip()
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run(code_text)
            self._apply_style(run)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        elif tag == "table":
            logging.info("Tableau détecté dans le contenu Markdown.")
            rows = elem.find_all("tr")
            if not rows:
                return

            header_cells = rows[0].find_all(["th", "td"])
            num_cols = len(header_cells)

            table = self.doc.add_table(rows=1, cols=num_cols)
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            for i, cell_elem in enumerate(header_cells):
                hdr_cells[i].text = ""
                paragraph = hdr_cells[i].paragraphs[0]
                self._add_inline(paragraph, cell_elem)

            for row_elem in rows[1:]:
                row_cells_elem = row_elem.find_all(["td", "th"])
                new_row_cells = table.add_row().cells
                for i, cell_elem in enumerate(row_cells_elem):
                    new_row_cells[i].text = ""
                    paragraph = new_row_cells[i].paragraphs[0]
                    self._add_inline(paragraph, cell_elem)
            logging.info("Tableau reconstruit dans le document DOCX.")
        else:
            text = elem.get_text(strip=True)
            if text:
                paragraph = self.doc.add_paragraph()
                self._add_inline(paragraph, elem)

    def add_markdown(self, text: str) -> None:
        """Convertit un texte Markdown et l'ajoute au document avec un fallback."""

        try:
            if not text:
                return

            md_converter = md.Markdown(extensions=["fenced_code", "tables"])
            html = md_converter.convert(text)

            soup = BeautifulSoup(html, "lxml")
            if soup.body:
                for elem in soup.body.find_all(recursive=False):
                    self._process_element(elem)
        except Exception as e:  # pragma: no cover - fallback branch
            warning_p = self.doc.add_paragraph()
            warning_run = warning_p.add_run(
                f"[Le formatage de ce bloc a échoué. Contenu original ci-dessous. Erreur : {e}]"
            )
            warning_run.font.italic = True
            warning_run.font.color.rgb = RGBColor(120, 120, 120)

            self.doc.add_paragraph(text)



def generer_export_docx_markdown(texte_markdown: str, styles_interface: Dict) -> io.BytesIO:
    """Génère un DOCX à partir d'un texte Markdown (cas de repli pour PDF/erreurs)."""
    document = Document()
    converter = MarkdownToDocxConverter(document, styles_interface)
    converter.add_markdown(texte_markdown)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output

