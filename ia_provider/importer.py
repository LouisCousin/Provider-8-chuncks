from __future__ import annotations

import logging
import re
from typing import Tuple, List, TypedDict

import docx
from docx.opc.exceptions import OpcError
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
import fitz  # PyMuPDF

# Configuration de la journalisation
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)


class BlocSemantique(TypedDict):
    niveau_titre: int
    contenu_markdown: str
    nombre_mots: int


def styles_contains_heading(document: docx.Document) -> bool:
    """Vérifie si au moins un paragraphe utilise un style de titre."""

    for para in document.paragraphs:
        style_name = para.style.name.lower()
        if "heading" in style_name or "titre" in style_name:
            return True
    return False


def _est_un_titre_probable(paragraph: Paragraph) -> int | None:
    """Détermine si un paragraphe ressemble à un titre.

    Retourne 1 ou 2 selon le niveau estimé, ou None s'il ne s'agit pas d'un titre.
    La détection est basée sur plusieurs critères :
    - brièveté et absence de ponctuation finale (rapides à vérifier)
    - tout le texte est en gras
    - numérotation éventuelle pour déterminer le niveau
    """

    text = paragraph.text.strip()
    if not text:
        return None

    # Vérifications rapides
    if len(text.split()) >= 15:
        return None
    if text.endswith('.'):
        return None

    runs_with_text = [run for run in paragraph.runs if run.text.strip()]
    if not runs_with_text:
        return None
    if not all(run.bold for run in runs_with_text):
        return None

    if re.match(r"^((\d+)|([A-Z])|([IVXLCDM]+))\.\s", text, re.IGNORECASE):
        return 1
    if re.match(r"^\d+(\.\d+)+\.?\s", text):
        return 2
    return 2

def _creer_blocs_semantiques(document: docx.Document, strategie: str) -> List[BlocSemantique]:
    """Transforme un document DOCX en blocs sémantiques."""

    blocs: List[BlocSemantique] = []
    contenu_actuel: List[str] = []
    niveau_actuel = 0

    def finaliser_bloc():
        nonlocal contenu_actuel, niveau_actuel
        if contenu_actuel:
            texte = "\n".join(contenu_actuel).strip()
            blocs.append(
                {
                    "niveau_titre": niveau_actuel,
                    "contenu_markdown": texte,
                    "nombre_mots": len(texte.split()),
                }
            )
            contenu_actuel = []
            niveau_actuel = 0

    for element in document.element.body:
        if isinstance(element, CT_P):
            paragraphe = Paragraph(element, document)
            style_name = (
                paragraphe.style.name
                if getattr(paragraphe, "style", None) and getattr(paragraphe.style, "name", None)
                else ""
            )
            style_lower = style_name.lower()
            style_normalized = style_lower.replace(" ", "").replace("-", "").replace("_", "")

            contenu = ""
            for run in paragraphe.runs:
                texte = run.text
                if not texte:
                    continue
                if run.bold and run.italic:
                    texte = f"***{texte}***"
                elif run.bold:
                    texte = f"**{texte}**"
                elif run.italic:
                    texte = f"*{texte}*"
                contenu += texte

            texte_stripped = contenu.strip()
            if not texte_stripped:
                contenu_actuel.append("")
                continue

            heading_level = None
            if strategie == "STYLE":
                for i in range(1, 5):
                    if f"heading{i}" in style_normalized or f"titre{i}" in style_normalized:
                        heading_level = i
                        break
            else:  # HEURISTIQUE
                heading_level = _est_un_titre_probable(paragraphe)

            if heading_level is not None:
                finaliser_bloc()
                contenu_actuel = ["#" * heading_level + f" {texte_stripped}"]
                niveau_actuel = heading_level
            elif "list bullet" in style_lower:
                contenu_actuel.append(f"* {texte_stripped}")
            else:
                contenu_actuel.append(texte_stripped)

        elif isinstance(element, CT_Tbl):
            table = Table(element, document)
            lignes_tableau: List[List[str]] = []
            for row in table.rows:
                cellules = []
                for cell in row.cells:
                    textes_cellule: List[str] = []
                    for para in cell.paragraphs:
                        contenu_para = ""
                        for run in para.runs:
                            texte = run.text
                            if not texte:
                                continue
                            if run.bold and run.italic:
                                texte = f"***{texte}***"
                            elif run.bold:
                                texte = f"**{texte}**"
                            elif run.italic:
                                texte = f"*{texte}*"
                            contenu_para += texte
                        textes_cellule.append(contenu_para or "")
                    cellules.append("<br>".join(textes_cellule))
                lignes_tableau.append(cellules)

            if lignes_tableau:
                entete = "| " + " | ".join(lignes_tableau[0]) + " |"
                separateur = "| " + " | ".join(["---"] * len(lignes_tableau[0])) + " |"
                contenu_actuel.append(entete)
                contenu_actuel.append(separateur)
                for ligne in lignes_tableau[1:]:
                    contenu_actuel.append("| " + " | ".join(ligne) + " |")
                contenu_actuel.append("")

    finaliser_bloc()
    return blocs


def _regrouper_blocs_en_chunks(blocs: List[BlocSemantique], mots_par_chunk: int) -> List[str]:
    """Regroupe les blocs sémantiques en chunks de taille cible."""
    if not blocs:
        return []

    chunks_finaux: List[str] = []
    chunk_actuel_contenu: List[str] = []
    chunk_actuel_mots = 0

    for bloc in blocs:
        if bloc["nombre_mots"] > mots_par_chunk:
            if chunk_actuel_contenu:
                chunks_finaux.append("\n".join(chunk_actuel_contenu))
            chunks_finaux.append(bloc["contenu_markdown"])
            chunk_actuel_contenu = []
            chunk_actuel_mots = 0
            continue

        if (chunk_actuel_mots + bloc["nombre_mots"]) > mots_par_chunk and chunk_actuel_mots > 0:
            chunks_finaux.append("\n".join(chunk_actuel_contenu))
            chunk_actuel_contenu = [bloc["contenu_markdown"]]
            chunk_actuel_mots = bloc["nombre_mots"]
        else:
            chunk_actuel_contenu.append(bloc["contenu_markdown"])
            chunk_actuel_mots += bloc["nombre_mots"]

    if chunk_actuel_contenu:
        chunks_finaux.append("\n".join(chunk_actuel_contenu))

    return chunks_finaux


def analyser_docx(file_stream) -> Tuple[List[str], str]:
    """Analyse un fichier DOCX et retourne les chunks générés et la stratégie."""

    try:
        file_stream.seek(0)
        document = docx.Document(file_stream)

        strategie = "STYLE" if styles_contains_heading(document) else "HEURISTIQUE"
        blocs = _creer_blocs_semantiques(document, strategie)
        MOTS_PAR_PAGE = 500
        TAILLE_CIBLE_PAGES = 15
        chunks = _regrouper_blocs_en_chunks(blocs, MOTS_PAR_PAGE * TAILLE_CIBLE_PAGES)
        return chunks, strategie
    except OpcError as e:
        logging.error(f"Fichier DOCX corrompu : {e}")
        return [], "HEURISTIQUE"
    except Exception as e:  # pragma: no cover - erreurs inattendues
        logging.error(f"Erreur inattendue sur DOCX : {e}", exc_info=True)
        return [], "HEURISTIQUE"


def analyser_pdf(file_stream) -> Tuple[List[str], str]:
    """Extrait le contenu textuel brut d'un PDF."""

    try:
        file_stream.seek(0)
        with fitz.open(stream=file_stream.read(), filetype="pdf") as doc:
            full_text = "".join(page.get_text() for page in doc)
        return [full_text], "IA"
    except Exception as e:  # pragma: no cover - erreurs inattendues
        logging.error(f"Erreur inattendue sur PDF : {e}", exc_info=True)
        return [""], "IA"


def analyser_document(fichier) -> Tuple[List[str], str]:
    """Analyse un fichier importé et choisit la méthode appropriée."""

    filename = fichier.name.lower()
    if filename.endswith(".docx"):
        return analyser_docx(fichier)
    if filename.endswith(".pdf"):
        return analyser_pdf(fichier)
    return [""], "STYLE"

